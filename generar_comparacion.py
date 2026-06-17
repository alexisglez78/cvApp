from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Estilos globales ────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Función de utilidad para encabezados de sección ─────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def colored_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_comparison_table(rows_data):
    """rows_data = list of (aspecto, original, localhost, impacto)"""
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ['ASPECTO', 'ORIGINAL (hospitalclinicanova.com)', 'LOCALHOST (:44373)', 'IMPACTO / PRIORIDAD']
    hdr_colors = ['1F4E79', '1F4E79', '1F4E79', '1F4E79']
    for i, (hdr, clr) in enumerate(zip(headers, hdr_colors)):
        cell = table.rows[0].cells[i]
        colored_cell(cell, clr)
        run = cell.paragraphs[0].add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    # Column widths
    widths = [Cm(3.5), Cm(6.5), Cm(6.5), Cm(3.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # Data rows
    impact_colors = {
        '🔴 ALTA': 'FFE0E0',
        '🟡 MEDIA': 'FFF9C4',
        '🟢 BAJA': 'E8F5E9',
    }
    for r_idx, (aspecto, original, localhost, impacto) in enumerate(rows_data, start=1):
        row = table.rows[r_idx]
        row.cells[0].paragraphs[0].add_run(aspecto).bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row.cells[1].paragraphs[0].add_run(original).font.size = Pt(9.5)
        row.cells[2].paragraphs[0].add_run(localhost).font.size = Pt(9.5)
        imp_run = row.cells[3].paragraphs[0].add_run(impacto)
        imp_run.bold = True
        imp_run.font.size = Pt(9.5)
        bg = impact_colors.get(impacto.split(' ')[0] + ' ' + impacto.split(' ')[1] if len(impacto.split()) > 1 else impacto, 'FFFFFF')
        for key, val in impact_colors.items():
            if key in impacto:
                colored_cell(row.cells[3], val)
                break
        # Zebra
        if r_idx % 2 == 0:
            for ci in [0, 1, 2]:
                colored_cell(row.cells[ci], 'F5F5F5')

    doc.add_paragraph()
    return table

# ════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('ANÁLISIS COMPARATIVO DE PÁGINAS WEB', 0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Hospital Clínica Nova — Original vs. Versión Local')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(31, 78, 121)

doc.add_paragraph()
info_table = doc.add_table(rows=3, cols=2)
info_table.style = 'Table Grid'
data = [
    ('URL Original', 'https://hospitalclinicanova.com/'),
    ('URL Local', 'https://localhost:44373/'),
    ('Fecha de análisis', '16 de junio de 2026'),
]
for i, (k, v) in enumerate(data):
    info_table.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
    info_table.rows[i].cells[1].paragraphs[0].add_run(v)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  1. RESUMEN EJECUTIVO
# ════════════════════════════════════════════════════════════════════════════
heading('1. RESUMEN EJECUTIVO', 1, (31, 78, 121))
doc.add_paragraph(
    'Se compararon visualmente y estructuralmente ambas versiones del sitio web de Hospital Clínica Nova. '
    'La versión local (localhost:44373) es una reimplementación del sitio original desarrollada en ASP.NET/Razor. '
    'Se identificaron 35+ diferencias en layout, tipografía, colores, espaciado, imágenes, animaciones y contenido textual. '
    'Las diferencias de mayor impacto visual están concentradas en el Hero, la navegación y la sección de Especialidades Médicas.'
)
add_divider()

# ════════════════════════════════════════════════════════════════════════════
#  2. NAVEGACIÓN / HEADER
# ════════════════════════════════════════════════════════════════════════════
heading('2. NAVEGACIÓN / HEADER', 1, (31, 78, 121))
add_comparison_table([
    (
        'Posición del logo',
        'Esquina superior izquierda, alineado a la izquierda con padding lateral',
        'Esquina superior izquierda — similar, pero con mayor margen inferior visible',
        '🟢 BAJA'
    ),
    (
        'Iconos del header',
        'Lupa (search) + icono hamburger menú — ambos a la derecha',
        'Icono hamburger al CENTRO, iconos de Facebook y YouTube a la derecha',
        '🔴 ALTA'
    ),
    (
        'Barra de búsqueda',
        'Buscador expandible oculto, activado con clic en la lupa',
        'Sin barra de búsqueda visible ni funcional',
        '🔴 ALTA'
    ),
    (
        'Redes sociales en header',
        'No aparecen en el header (sólo en footer)',
        'Facebook y YouTube visibles directamente en el header',
        '🟡 MEDIA'
    ),
    (
        'Color de fondo del header',
        'Blanco (#FFFFFF) con sombra suave',
        'Blanco (#FFFFFF) — similar pero sin sombra perceptible',
        '🟢 BAJA'
    ),
    (
        'Menú lateral (sidebar)',
        'Menú lateral deslizable con lista completa de navegación',
        'Menú hamburger desplegable (Bootstrap collapse)',
        '🟡 MEDIA'
    ),
    (
        'Altura del header',
        '~60px de altura, compacto',
        '~70px de altura, ligeramente más alto',
        '🟢 BAJA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  3. HERO / BANNER PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════
heading('3. HERO / BANNER PRINCIPAL', 1, (31, 78, 121))
add_comparison_table([
    (
        'Fondo del Hero',
        'Fondo 100% naranja/ámbar (#F5A623 aprox.) — NO hay imagen de fondo visible en primer frame',
        'Imagen de médico/dentista trabajando SUPERPUESTA sobre el bloque naranja, en esquina superior izquierda',
        '🔴 ALTA'
    ),
    (
        'Layout del Hero',
        'Texto centrado a la izquierda sobre fondo naranja puro. Sin imagen.',
        'Imagen rectangular con bordes redondeados flotando encima del bloque naranja de texto',
        '🔴 ALTA'
    ),
    (
        'Título H1',
        '"Hospital Clínica Nova" — tamaño aprox. 38-42px, fuente blanca, peso bold, alineado izquierda',
        '"Hospital Clínica Nova" — tamaño aprox. 28-32px, centrado dentro del bloque naranja',
        '🔴 ALTA'
    ),
    (
        'Subtítulo',
        'Etiqueta H5, ~18px, blanco, alineado izquierda. Texto completo en 2 líneas.',
        'Párrafo <p>, ~15px, blanco, centrado. Texto en 3 líneas.',
        '🔴 ALTA'
    ),
    (
        'Botón "Conoce más"',
        'Botón rectangular con bordes rectos, fondo azul marino (#1A3C6E approx.), texto blanco, padding ~12px 24px',
        'Botón con fondo azul más claro (#2196F3 approx.), texto blanco, bordes ligeramente redondeados',
        '🟡 MEDIA'
    ),
    (
        'Color de fondo naranja',
        'Naranja ámbar: aproximadamente #F5A51F / #F0A500',
        'Naranja similar pero ligeramente más saturado/oscuro — aprox. #E89A00',
        '🟢 BAJA'
    ),
    (
        'Slider/Carrusel',
        'Carrusel de imágenes/banners superpuesto en la parte derecha del hero, con 4 slides',
        'Carrusel con botones Previous/Next visibles, 4 imágenes con href="javascript:void(0)" (no navegan)',
        '🟡 MEDIA'
    ),
    (
        'Altura del Hero',
        'Hero ocupa ~80% de la altura del viewport (pantalla completa)',
        'Hero mucho más pequeño — aprox. 50-60% del viewport',
        '🔴 ALTA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  4. SECCIÓN "NOVA TE CONTESTA"
# ════════════════════════════════════════════════════════════════════════════
heading('4. SECCIÓN "NOVA TE CONTESTA"', 1, (31, 78, 121))
add_comparison_table([
    (
        'Imagen destacada',
        'Imagen grande de celular con la app "Nova Contesta" a la izquierda de la sección',
        'Sin imagen del celular — solo el carousel de videos y texto',
        '🔴 ALTA'
    ),
    (
        'Título de la sección',
        'Heading H4 "Nova te contesta" — mayor jerarquía tipográfica, más prominente',
        'Párrafo <p> simple — sin jerarquía de heading, visualmente más pequeño',
        '🟡 MEDIA'
    ),
    (
        'Videos de YouTube',
        'Carousel de 3 videos con thumbnail visible + descripción a la derecha de cada video',
        'Carousel con 1 video visible a la vez (mismo carousel Bootstrap) — sin imagen extra',
        '🟡 MEDIA'
    ),
    (
        'Descripción del video',
        'Heading H3 + párrafo descripción al lado derecho del video en desktop',
        'H3 y párrafo debajo del video — layout apilado verticalmente',
        '🟡 MEDIA'
    ),
    (
        'Texto "Clínica de Salud Mental"',
        '"Clínica de Salud Mental y Neurociencias" (con acento en Clínica y plural en Neurociencias)',
        '"Clinica de Salud Mental y Neurociencia" — SIN acento en "Clínica", sin "s" en "Neurociencias"',
        '🟡 MEDIA'
    ),
    (
        'Controles del carousel',
        'Flechas left/right integradas con diseño personalizado del tema',
        'Botones Bootstrap estándar "Previous" / "Next" con texto visible',
        '🟢 BAJA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  5. SECCIÓN ESPECIALIDADES MÉDICAS
# ════════════════════════════════════════════════════════════════════════════
heading('5. SECCIÓN ESPECIALIDADES MÉDICAS', 1, (31, 78, 121))
add_comparison_table([
    (
        'Presentación de especialidades',
        'Lista de texto simple en 2 columnas — SIN iconos ni imágenes',
        'Grid con IMAGEN + texto para cada especialidad — mucho más visual',
        '🔴 ALTA'
    ),
    (
        'Imágenes por especialidad',
        'No hay imágenes por especialidad',
        'Cada especialidad tiene su propia imagen/icono representativo',
        '🔴 ALTA'
    ),
    (
        'Acentos ortográficos',
        'Cardiología, Gastroenterología, Hematología, Infectología, Neumología, Oncología, Perinatología, Psiquiatría, Radiología, Reumatología (CORRECTOS)',
        'Cardiologia, Gastroenterologia, Hematologia, Infectologia, Neumologia, Oncologia, Perinatologia, Psiquiatria, Radiologia, Reumatologia (SIN ACENTOS)',
        '🔴 ALTA'
    ),
    (
        'Botón "Descubre más"',
        'Botón centrado debajo de la lista, diseño consistente con el tema',
        'Botón presente pero con estilo Bootstrap genérico',
        '🟢 BAJA'
    ),
    (
        'Layout de columnas',
        '2 columnas de 6 items cada una (texto puro)',
        'Grid 6 columnas por fila con imagen + texto — diferente distribución',
        '🔴 ALTA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  6. SECCIÓN COVID-19
# ════════════════════════════════════════════════════════════════════════════
heading('6. SECCIÓN COVID-19', 1, (31, 78, 121))
add_comparison_table([
    (
        'Título de sección',
        '"RESPUESTA ANTE EL COVID-19" — heading H3 unificado, con word-break tipográfico intencional',
        '"RESPUESTA" y "ANTE EL COVID-19" en dos párrafos <p> separados — falla visual',
        '🔴 ALTA'
    ),
    (
        'Estadísticas numéricas',
        'H5 con animación/highlight en los números grandes ("3.7 millones de USD")',
        'Párrafo <p> simple sin jerarquía ni estilo especial en los números',
        '🟡 MEDIA'
    ),
    (
        'Imagen de fondo',
        'Imagen de fondo oscura con overlay para contraste del texto',
        'Sin imagen de fondo — fondo liso (blanco o gris claro)',
        '🔴 ALTA'
    ),
    (
        'Links subcategorías COVID',
        '3 links con imagen thumbnail: "Hospital Comunitario", "TEMI", "Investigaciones"',
        'Sin estos links/thumbnails — sección incompleta',
        '🔴 ALTA'
    ),
    (
        'Botón "Mira como lo hizo"',
        'Enlaza a /covid-19',
        'Enlaza a /investigacion-y-ensenanza-covid-19/ — URL diferente',
        '🟡 MEDIA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  7. SECCIÓN "ACERCA DE"
# ════════════════════════════════════════════════════════════════════════════
heading('7. SECCIÓN "ACERCA DE HOSPITAL CLÍNICA NOVA"', 1, (31, 78, 121))
add_comparison_table([
    (
        'Jerarquía de títulos',
        'H3 para "Acerca de", H6 para "NUESTRA MISIÓN" y "VISIÓN" — estructura semántica correcta',
        'Párrafos <p> para todos los textos — sin jerarquía HTML semántica',
        '🔴 ALTA'
    ),
    (
        'Imagen de fondo/lateral',
        'Imagen de fondo del hospital visible en esta sección',
        'Sin imagen — solo texto sobre fondo liso',
        '🔴 ALTA'
    ),
    (
        'Botón "Explora más"',
        'Enlaza a https://hospitalclinicanova.com/nosotros',
        'Enlaza a /nosotros/ — relativo (funciona igual en contexto local)',
        '🟢 BAJA'
    ),
    (
        'Sub-sección calidad + directorio',
        'H4 con texto + link "Ver más" cada uno — layout en 2 columnas side by side',
        'Layout en 2 columnas también, pero con párrafos <p> en lugar de H4 — tipografía más pequeña',
        '🟡 MEDIA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  8. SECCIÓN NOTICIAS
# ════════════════════════════════════════════════════════════════════════════
heading('8. SECCIÓN "HISTORIAS PARA COMPARTIR" / NOTICIAS', 1, (31, 78, 121))
add_comparison_table([
    (
        'Título de sección',
        '"Historias para compartir" — heading H3',
        'Sin título de sección — las noticias aparecen directamente sin encabezado',
        '🔴 ALTA'
    ),
    (
        'Primera noticia',
        'Card con imagen + título + link directo al artículo completo',
        'Link de texto puro con párrafo — sin imagen de portada en la primera noticia',
        '🔴 ALTA'
    ),
    (
        'Segunda/Tercera noticia',
        'Cards con imagen thumbnail + título',
        'Links con imagen + párrafo — layout similar pero con diferencia tipográfica',
        '🟡 MEDIA'
    ),
    (
        'Texto con acentos — noticia 1',
        '"¡Refuerza las medidas de prevención de COVID-19!" (con acento en "prevención")',
        '"¡Refuerza las medidas de prevencion de COVID-19!" (SIN acento en "prevención")',
        '🟡 MEDIA'
    ),
    (
        'Título noticia 2',
        '"¿Sabes cómo hacer la maniobra de Heimlich?"',
        '"¿Sabes cómo hacer la maniobra Heimlich?" — falta "de"',
        '🟢 BAJA'
    ),
    (
        'URLs de noticias',
        'URLs semánticas: /refuerza-las-medidas..., /sabes-como-hacer...',
        'URLs genéricas: /noticias/noticias_pag_1, /noticias/noticias_pag_2, /noticias/noticias_pag_3',
        '🟡 MEDIA'
    ),
    (
        'Botón "Entérate de más"',
        'Presente, enlaza a /noticias',
        'Presente, enlaza a /noticias/ — igual en funcionalidad',
        '🟢 BAJA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  9. SECCIÓN DERECHOHABIENTES
# ════════════════════════════════════════════════════════════════════════════
heading('9. SECCIÓN "DERECHOHABIENTES NOVA"', 1, (31, 78, 121))
add_comparison_table([
    (
        'Icono portal servicios',
        'Dos iconos/imágenes diferenciados para cada acceso',
        'Texto con ">" como prefijo en lugar de iconos — aspecto no pulido',
        '🔴 ALTA'
    ),
    (
        'Botón "Ir a portal"',
        'Botón con diseño del tema, enlaza a novaservicios.com.mx',
        'Botón Bootstrap estándar, mismo enlace',
        '🟢 BAJA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  10. FOOTER
# ════════════════════════════════════════════════════════════════════════════
heading('10. FOOTER', 1, (31, 78, 121))
add_comparison_table([
    (
        'Estructura del footer',
        'Footer con: logo, eslogan, redes sociales, navegación, copyright, términos y privacidad',
        'No hay footer visible en la captura del localhost — posiblemente ausente o incompleto',
        '🔴 ALTA'
    ),
    (
        'Redes sociales en footer',
        'Facebook + YouTube con iconos del tema',
        'Solo en el header — no en footer',
        '🟡 MEDIA'
    ),
    (
        'Copyright',
        '"© 2026 Hospital Clínica Nova."',
        'No detectado en la versión local',
        '🟢 BAJA'
    ),
    (
        'Links legales',
        'Términos y condiciones + Aviso de privacidad',
        'No detectados',
        '🟡 MEDIA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  11. TIPOGRAFÍA
# ════════════════════════════════════════════════════════════════════════════
heading('11. TIPOGRAFÍA', 1, (31, 78, 121))
add_comparison_table([
    (
        'Fuente principal',
        'Fuente personalizada del tema WordPress (sans-serif, posiblemente "Montserrat" o "Open Sans")',
        'Bootstrap 5 default — "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif',
        '🟡 MEDIA'
    ),
    (
        'Tamaño H1 Hero',
        '~38-42px, negrita, peso 700-800',
        '~28-32px, negrita — perceptiblemente más pequeño',
        '🔴 ALTA'
    ),
    (
        'Tamaño H5 subtítulo Hero',
        '~18-20px, peso 400-500',
        '~15-16px párrafo normal — menor jerarquía',
        '🔴 ALTA'
    ),
    (
        'Uso de headings HTML',
        'H1, H3, H4, H5, H6 usados correctamente según jerarquía semántica',
        'La mayoría de títulos son párrafos <p> — sin jerarquía semántica',
        '🔴 ALTA'
    ),
    (
        'Acentos en textos',
        'Todos los textos con acentos correctos (tildes en español)',
        'Múltiples palabras sin acento: prevención→prevencion, Cardiología→Cardiologia, etc.',
        '🔴 ALTA'
    ),
    (
        'Interlineado (line-height)',
        'Mayor interlineado (~1.6-1.8) — textos más legibles y aireados',
        'Bootstrap default (~1.5) — más compacto',
        '🟢 BAJA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  12. COLORES
# ════════════════════════════════════════════════════════════════════════════
heading('12. PALETA DE COLORES', 1, (31, 78, 121))
add_comparison_table([
    (
        'Color principal (naranja)',
        '#F5A623 aproximado — naranja ámbar cálido',
        'Similar pero ligeramente más oscuro/saturado',
        '🟢 BAJA'
    ),
    (
        'Color botón primario',
        'Azul marino oscuro (~#1A3C6E) — alto contraste',
        'Azul Bootstrap (#0d6efd) o similar — más brillante',
        '🟡 MEDIA'
    ),
    (
        'Fondo sección COVID',
        'Fondo oscuro (imagen + overlay negro semitransparente)',
        'Fondo liso claro (sin imagen de fondo)',
        '🔴 ALTA'
    ),
    (
        'Texto sobre naranja',
        'Blanco puro (#FFFFFF) — contraste alto (AA/AAA)',
        'Blanco puro (#FFFFFF) — igual',
        '🟢 BAJA'
    ),
    (
        'Color de links en footer',
        'Blanco o gris claro con hover naranja',
        'Footer no implementado',
        '🟡 MEDIA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  13. ESPACIADO Y LAYOUT
# ════════════════════════════════════════════════════════════════════════════
heading('13. ESPACIADO Y LAYOUT', 1, (31, 78, 121))
add_comparison_table([
    (
        'Padding sección Hero',
        '~80-100px vertical, ~5% horizontal — muy generoso',
        '~40-50px vertical — mucho más compacto',
        '🔴 ALTA'
    ),
    (
        'Ancho máximo del contenido',
        'Container con max-width ~1200px + padding lateral',
        'Container Bootstrap estándar (max ~1140px) — similar',
        '🟢 BAJA'
    ),
    (
        'Espaciado entre secciones',
        '~60-80px entre secciones (margin-bottom generoso)',
        '~30-40px entre secciones — secciones más "pegadas"',
        '🟡 MEDIA'
    ),
    (
        'Padding interno de cards',
        '~24-32px en cards y bloques internos',
        '~16-20px — cards más compactas',
        '🟡 MEDIA'
    ),
    (
        'Grid de especialidades',
        '2 columnas de lista (texto)',
        'Grid de 6 columnas (imagen + texto) — diferente sistema',
        '🔴 ALTA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  14. ANIMACIONES E INTERACTIVIDAD
# ════════════════════════════════════════════════════════════════════════════
heading('14. ANIMACIONES E INTERACTIVIDAD', 1, (31, 78, 121))
add_comparison_table([
    (
        'Carrusel hero / slider banners',
        'Slider automático con transición fade/slide suave, controles personalizados',
        'Carrusel Bootstrap con botones Previous/Next estándar (texto visible)',
        '🟡 MEDIA'
    ),
    (
        'Transición de hover en links',
        'Efectos hover personalizados (underline animado, cambio de color suave)',
        'Hover Bootstrap estándar — menos elaborado',
        '🟢 BAJA'
    ),
    (
        'Animación en estadísticas COVID',
        'Posiblemente contador animado en los números (3.7M, 350K, 31.4K)',
        'Sin animación de contador — texto estático',
        '🟡 MEDIA'
    ),
    (
        'Scroll behavior',
        'Smooth scroll al hacer clic en "Skip to main content"',
        'Sin enlace "Skip to main content"',
        '🟢 BAJA'
    ),
    (
        'Lazy loading de imágenes',
        'Imágenes con lazy load (WordPress optimizado)',
        'Imágenes cargadas directamente sin lazy load evidente',
        '🟢 BAJA'
    ),
    (
        'Menú lateral',
        'Sidebar menu deslizable desde la derecha con overlay oscuro',
        'Menú colapsable Bootstrap — sin overlay ni slide animado',
        '🟡 MEDIA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  15. IMÁGENES
# ════════════════════════════════════════════════════════════════════════════
heading('15. IMÁGENES', 1, (31, 78, 121))
add_comparison_table([
    (
        'Imagen en Hero',
        'Sin imagen propia en hero (solo fondo naranja + texto)',
        'Imagen de procedimiento dental/médico flotando sobre el bloque naranja',
        '🔴 ALTA'
    ),
    (
        'Imagen "Nova te contesta"',
        'Imagen de celular con app Nova Contesta (muy relevante para la sección)',
        'Sin esta imagen — sección visualmente incompleta',
        '🔴 ALTA'
    ),
    (
        'Thumbnails COVID',
        '3 imágenes de thumbnail para Hospital Comunitario, TEMI, Investigaciones',
        'Sin estas imágenes',
        '🔴 ALTA'
    ),
    (
        'Imagen primera noticia',
        'Card con imagen de portada de la noticia',
        'Sin imagen en la primera noticia (solo enlace de texto)',
        '🟡 MEDIA'
    ),
    (
        'Iconos especialidades',
        'Sin iconos/imágenes por especialidad (solo texto)',
        'Imágenes representativas por especialidad — VENTAJA del localhost',
        '🟢 BAJA'
    ),
    (
        'Imagen sección "Acerca de"',
        'Imagen del hospital o equipo médico como fondo/lateral',
        'Sin imagen — solo texto',
        '🔴 ALTA'
    ),
    (
        'Formato de imágenes',
        'WebP (.webp) para mejor rendimiento y calidad',
        'Formato desconocido (imágenes servidas desde servidor local)',
        '🟢 BAJA'
    ),
])

# ════════════════════════════════════════════════════════════════════════════
#  16. RESUMEN PRIORIDADES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading('16. RESUMEN DE PRIORIDADES PARA CORRECCIÓN', 1, (31, 78, 121))
doc.add_paragraph('Lista ordenada de correcciones por impacto visual:')

prioridades = [
    ('🔴 ALTA PRIORIDAD (corregir primero)', [
        '1. Hero: eliminar imagen de médico flotante — el hero original es solo fondo naranja con texto',
        '2. Hero: aumentar tamaño del H1 a ~38-42px y aumentar el padding/altura del hero',
        '3. Navegación: quitar redes sociales del header, agregar icono de búsqueda + menú a la derecha',
        '4. Implementar barra de búsqueda expandible',
        '5. Título COVID: unificar en un solo heading H3 en lugar de dos párrafos',
        '6. Añadir imagen de fondo oscura en la sección COVID',
        '7. Agregar los 3 thumbnails/links de subcategorías COVID (Hospital Comunitario, TEMI, Investigaciones)',
        '8. Sección "Nova te contesta": añadir imagen del celular con la app',
        '9. Sección "Acerca de": añadir imagen del hospital y usar headings HTML correctos',
        '10. Especialidades: cambiar grid de imágenes a lista de texto (para coincidir con original)',
        '11. Footer: implementar footer completo con logo, navegación, redes, copyright, links legales',
        '12. ACENTOS: corregir todos los textos sin tilde (Cardiología, prevención, Clínica, etc.)',
        '13. Jerarquía HTML: reemplazar <p> por headings H3/H4/H5/H6 donde corresponda',
    ]),
    ('🟡 MEDIA PRIORIDAD', [
        '14. Botón "Conoce más": ajustar color a azul marino oscuro más cercano al original',
        '15. Sección "Nova te contesta": título debe ser H4, no párrafo <p>',
        '16. Carousel de videos: mejorar controles (flechas personalizadas vs texto "Previous/Next"',
        '17. Espaciado entre secciones: aumentar margin-bottom a ~60-80px',
        '18. Padding del Hero: aumentar a ~80-100px vertical',
        '19. Footer: agregar redes sociales en footer (no solo en header)',
        '20. URLs de noticias: usar URLs semánticas en lugar de /noticias/noticias_pag_N',
    ]),
    ('🟢 BAJA PRIORIDAD', [
        '21. Fuente: considerar importar Montserrat u Open Sans para coincidir con el original',
        '22. Interlineado: aumentar line-height de 1.5 a ~1.6-1.8',
        '23. Altura del header: reducir de ~70px a ~60px',
        '24. Botón "Ir a portal": ajustar estilo del botón',
        '25. Añadir "Skip to main content" para accesibilidad',
    ]),
]

for section_title, items in prioridades:
    p = doc.add_paragraph()
    r = p.add_run(section_title)
    r.bold = True
    r.font.size = Pt(12)
    for item in items:
        bp = doc.add_paragraph(item, style='List Bullet')
        bp.runs[0].font.size = Pt(10)

# Fin del documento
doc.add_page_break()
heading('17. NOTAS FINALES', 1, (31, 78, 121))
doc.add_paragraph(
    'Este análisis se realizó comparando la versión en producción de hospitalclinicanova.com con la versión '
    'local en https://localhost:44373/ mediante captura de pantalla, análisis del DOM (árbol de accesibilidad) '
    'y extracción del contenido textual de ambas páginas. '
    'El documento cubre: navegación, hero, todas las secciones principales, tipografía, colores, espaciado, '
    'imágenes y animaciones. Total de diferencias identificadas: 35+.'
)

output_path = '/Users/pgoncasa/Documents/Comparacion_Nova_Original_vs_Local.docx'
doc.save(output_path)
print(f'✅ Documento generado: {output_path}')
